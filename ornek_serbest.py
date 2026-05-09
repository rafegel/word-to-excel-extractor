"""Serbest metin formatında örnek Word dosyası oluşturur."""
from docx import Document

doc = Document()
doc.add_heading("Yıllık Personel Değerlendirme Raporu — 2024", 0)

doc.add_paragraph(
    "Bu rapor, şirketimizin 2024 yılı personel performans değerlendirmesini kapsamaktadır. "
    "İnsan Kaynakları birimi tarafından hazırlanan bu belgede çalışanlara ait çeşitli bilgiler yer almaktadır."
)

doc.add_heading("Değerlendirme Sonuçları", 1)

doc.add_paragraph(
    "Yazılım geliştirme ekibinden Ahmet Yılmaz bu yıl olağanüstü bir performans sergilemiştir. "
    "Kıdemli Yazılım Geliştirici pozisyonunda görev yapan Ahmet, mevcut maaşı olan 45.000 TL'ye ek olarak "
    "yüzde on oranında zam almaya hak kazanmıştır. Kendisiyle ahmet@sirket.com adresinden iletişime geçilebilir."
)

doc.add_paragraph(
    "Aynı ekipten Zeynep Kaya, Backend Developer olarak yılı başarıyla tamamlamıştır. "
    "38.500 TL maaşla çalışan Zeynep'in performansı beklentileri karşılamaktadır. "
    "İletişim: zeynep@sirket.com"
)

doc.add_paragraph(
    "Frontend geliştirme tarafında Mert Demir dikkat çekici çalışmalar ortaya koymuştur. "
    "35.000 TL maaşla istihdam edilen Mert, müşteri arayüzlerinde önemli iyileştirmeler gerçekleştirmiştir. "
    "Eposta adresi: mert@sirket.com"
)

doc.add_heading("Pazarlama Ekibi", 1)

doc.add_paragraph(
    "Pazarlama departmanı bu yıl rekor satış rakamlarına ulaşmıştır. Departman müdürü Fatma Çelik, "
    "55.000 TL brüt maaşıyla ekibi başarıyla yönetmiştir. fatma@sirket.com adresinden ulaşılabilir. "
    "Fatma Hanım'ın liderliğinde ekip, hedeflerin yüzde yüz on beşine ulaşmıştır."
)

doc.add_paragraph(
    "Sosyal medya yönetimini üstlenen Can Öztürk, 28.000 TL maaşla şirketin dijital varlığını "
    "güçlendirmiştir. Instagram ve LinkedIn hesaplarında takipçi sayısını ikiye katlamıştır. "
    "can@sirket.com üzerinden kendisine ulaşılabilmektedir."
)

doc.add_paragraph(
    "İçerik stratejisi alanında Ayşe Şahin, 32.750 TL maaşla özgün içerikler üretmeye devam etmektedir. "
    "Blog ve e-posta kampanyalarında dönüşüm oranlarını artıran Ayşe'ye ayse@sirket.com adresinden ulaşılabilir."
)

doc.add_heading("Muhasebe ve Finans", 1)

doc.add_paragraph(
    "Finans departmanında Kemal Aydın, Mali Müşavir unvanıyla 60.000 TL maaş almaktadır. "
    "Yıl boyunca şirketin mali tablolarını titizlikle yönetmiştir. kemal@sirket.com"
)

doc.add_paragraph(
    "Muhasebe uzmanı Derya Bulut ise 34.000 TL maaşla fatura ve ödeme süreçlerini yürütmektedir. "
    "derya@sirket.com adresine yazılabilir."
)

doc.add_heading("Genel Değerlendirme", 1)
doc.add_paragraph(
    "2024 yılında toplam 8 çalışan değerlendirmeye alınmıştır. Şirket genelinde ortalama maaş artışı "
    "yüzde sekiz olarak gerçekleşmiştir. Bir sonraki değerlendirme Haziran 2025'te yapılacaktır. "
    "Sorularınız için ik@sirket.com adresine başvurabilirsiniz."
)

doc.save("ornek_serbest_metin.docx")
print("✅ ornek_serbest_metin.docx oluşturuldu.")
