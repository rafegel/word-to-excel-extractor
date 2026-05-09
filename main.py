#!/usr/bin/env python3
"""
Word/PDF → Excel Veri Çıkarıcı
Tkinter GUI + Watchdog klasör izleme + Groq AI + PDF desteği
"""

import sys
import json
import re
import threading
from pathlib import Path
from datetime import datetime

import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, messagebox

from docx import Document
from groq import Groq
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

try:
    import PyPDF2
    PDF_DESTEKLI = True
except ImportError:
    PDF_DESTEKLI = False


# ─── Dosya Okuma ─────────────────────────────────────────────────────────────

def word_oku(dosya_yolu: str) -> str:
    doc = Document(dosya_yolu)
    satirlar = []
    for paragraf in doc.paragraphs:
        if paragraf.text.strip():
            satirlar.append(paragraf.text.strip())
    for tablo in doc.tables:
        for satir in tablo.rows:
            hucreler = [h.text.strip() for h in satir.cells if h.text.strip()]
            if hucreler:
                satirlar.append(" | ".join(hucreler))
    return "\n".join(satirlar)


def pdf_oku(dosya_yolu: str) -> str:
    if not PDF_DESTEKLI:
        return ""
    with open(dosya_yolu, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        return "\n".join(
            page.extract_text() for page in reader.pages if page.extract_text()
        )


def dosya_oku(dosya_yolu: str) -> str:
    yol = Path(dosya_yolu)
    if yol.suffix.lower() == ".docx":
        return word_oku(dosya_yolu)
    elif yol.suffix.lower() == ".pdf":
        return pdf_oku(dosya_yolu)
    return ""


# ─── Groq & Excel ────────────────────────────────────────────────────────────

def groq_ile_cikart(metin: str, istek: str, api_key: str) -> dict:
    client = Groq(api_key=api_key)
    sistem = """Sen bir veri çıkarma asistanısın. Metinden istenen bilgileri çıkarıp SADECE geçerli JSON döndür.
Format: {"basliklar": ["Sütun1", ...], "satirlar": [["değer1", ...], ...]}
Kural: Sadece JSON, Türkçe karakterleri koru, veri yoksa boş liste döndür."""

    yanit = client.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": sistem},
            {"role": "user", "content": f"Şu bilgileri çıkar: {istek}\n\nMETİN:\n{metin[:8000]}"}
        ],
        temperature=0.1,
        max_tokens=4000
    )
    ham = yanit.choices[0].message.content.strip()
    ham = re.sub(r'^```json\s*', '', ham)
    ham = re.sub(r'\s*```$', '', ham)
    return json.loads(ham)


def excel_olustur(veri: dict, cikti_yolu: str):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Çıkarılan Veri"

    baslik_font = Font(name="Arial", bold=True, color="FFFFFF", size=11)
    baslik_fill = PatternFill("solid", start_color="2E75B6")
    baslik_hizalama = Alignment(horizontal="center", vertical="center", wrap_text=True)
    veri_font = Font(name="Arial", size=10)
    cift_fill = PatternFill("solid", start_color="D6E4F0")
    tek_fill = PatternFill("solid", start_color="FFFFFF")
    veri_hizalama = Alignment(vertical="center", wrap_text=True)
    ince = Side(style="thin", color="BDD7EE")
    sinir = Border(left=ince, right=ince, top=ince, bottom=ince)

    basliklar = veri.get("basliklar", [])
    satirlar = veri.get("satirlar", [])

    if not basliklar:
        ws["A1"] = "Veri bulunamadı."
        wb.save(cikti_yolu)
        return 0, 0

    for col_idx, baslik in enumerate(basliklar, 1):
        h = ws.cell(row=1, column=col_idx, value=baslik)
        h.font = baslik_font
        h.fill = baslik_fill
        h.alignment = baslik_hizalama
        h.border = sinir
    ws.row_dimensions[1].height = 30

    for row_idx, satir in enumerate(satirlar, 2):
        fill = cift_fill if row_idx % 2 == 0 else tek_fill
        for col_idx, deger in enumerate(satir, 1):
            h = ws.cell(row=row_idx, column=col_idx, value=deger)
            h.font = veri_font
            h.fill = fill
            h.alignment = veri_hizalama
            h.border = sinir
        ws.row_dimensions[row_idx].height = 20

    for col_idx, baslik in enumerate(basliklar, 1):
        col_letter = openpyxl.utils.get_column_letter(col_idx)
        max_uzunluk = len(str(baslik))
        for satir in satirlar:
            if col_idx - 1 < len(satir):
                max_uzunluk = max(max_uzunluk, len(str(satir[col_idx - 1])))
        ws.column_dimensions[col_letter].width = min(max_uzunluk + 4, 50)

    ws.freeze_panes = "A2"

    ws_ozet = wb.create_sheet("Özet")
    ws_ozet["A1"] = "Toplam Satır:"
    ws_ozet["B1"] = len(satirlar)
    ws_ozet["A2"] = "Toplam Sütun:"
    ws_ozet["B2"] = len(basliklar)
    for hucre in ["A1", "A2"]:
        ws_ozet[hucre].font = Font(name="Arial", bold=True)

    wb.save(cikti_yolu)
    return len(satirlar), len(basliklar)


# ─── Watchdog Handler ─────────────────────────────────────────────────────────

class DosyaHandler(FileSystemEventHandler):
    def __init__(self, uygulama):
        self.uygulama = uygulama
        self.islenen = set()

    def on_created(self, event):
        if event.is_directory:
            return
        yol = Path(event.src_path)
        if yol.suffix.lower() in (".docx", ".pdf") and str(yol) not in self.islenen:
            self.islenen.add(str(yol))
            self.uygulama.dosya_isle(str(yol))


# ─── GUI ─────────────────────────────────────────────────────────────────────

class Uygulama:
    def __init__(self, root):
        self.root = root
        self.root.title("📄➡️📊 Word/PDF → Excel Çıkarıcı")
        self.root.geometry("700x600")
        self.root.configure(bg="#1e1e2e")
        self.root.resizable(True, True)

        self.observer = None
        self.izleniyor = False

        self._arayuz_olustur()

    def _arayuz_olustur(self):
        # Renkler
        BG = "#1e1e2e"
        SURFACE = "#2a2a3e"
        ACCENT = "#7c6af7"
        TEXT = "#e8e8f0"
        MUTED = "#6b6b80"
        GREEN = "#4ade80"
        RED = "#f87171"

        self.root.configure(bg=BG)

        # Başlık
        baslik = tk.Label(self.root, text="📄➡️📊 Word/PDF → Excel Çıkarıcı",
                         font=("Segoe UI", 16, "bold"), bg=BG, fg=TEXT)
        baslik.pack(pady=(20, 5))

        alt = tk.Label(self.root, text="Klasörü izle, dosya gelince otomatik Excel üret",
                      font=("Segoe UI", 10), bg=BG, fg=MUTED)
        alt.pack(pady=(0, 15))

        # Ana frame
        ana = tk.Frame(self.root, bg=BG, padx=20)
        ana.pack(fill="both", expand=True)

        # API Key
        tk.Label(ana, text="🔑 Groq API Anahtarı", font=("Segoe UI", 10, "bold"),
                bg=BG, fg=TEXT).pack(anchor="w")
        self.api_key = tk.Entry(ana, show="*", font=("Segoe UI", 10),
                               bg=SURFACE, fg=TEXT, insertbackground=TEXT,
                               relief="flat", bd=5)
        self.api_key.pack(fill="x", pady=(3, 12))

        # İzlenecek klasör
        tk.Label(ana, text="📁 İzlenecek Klasör", font=("Segoe UI", 10, "bold"),
                bg=BG, fg=TEXT).pack(anchor="w")
        klasor_frame = tk.Frame(ana, bg=BG)
        klasor_frame.pack(fill="x", pady=(3, 12))

        self.klasor_yolu = tk.StringVar()
        tk.Entry(klasor_frame, textvariable=self.klasor_yolu, font=("Segoe UI", 10),
                bg=SURFACE, fg=TEXT, insertbackground=TEXT, relief="flat", bd=5).pack(
                side="left", fill="x", expand=True)
        tk.Button(klasor_frame, text="Seç", font=("Segoe UI", 9, "bold"),
                 bg=ACCENT, fg="white", relief="flat", padx=10,
                 command=self._klasor_sec).pack(side="right", padx=(8, 0))

        # Veri isteği
        tk.Label(ana, text="🤖 Ne çıkartılsın? (doğal dille yaz)",
                font=("Segoe UI", 10, "bold"), bg=BG, fg=TEXT).pack(anchor="w")
        self.istek = tk.Entry(ana, font=("Segoe UI", 10),
                             bg=SURFACE, fg=TEXT, insertbackground=TEXT,
                             relief="flat", bd=5)
        self.istek.insert(0, "tüm veriler")
        self.istek.pack(fill="x", pady=(3, 12))

        # Butonlar
        btn_frame = tk.Frame(ana, bg=BG)
        btn_frame.pack(fill="x", pady=(5, 12))

        self.btn_izle = tk.Button(btn_frame, text="▶ İzlemeyi Başlat",
                                 font=("Segoe UI", 10, "bold"),
                                 bg=GREEN, fg="#000", relief="flat", padx=15, pady=8,
                                 command=self._izleme_toggle)
        self.btn_izle.pack(side="left")

        tk.Button(btn_frame, text="📂 Dosya Seç & İşle",
                 font=("Segoe UI", 10, "bold"),
                 bg=ACCENT, fg="white", relief="flat", padx=15, pady=8,
                 command=self._manuel_isle).pack(side="left", padx=(10, 0))

        tk.Button(btn_frame, text="🗑 Temizle",
                 font=("Segoe UI", 10, "bold"),
                 bg=SURFACE, fg=MUTED, relief="flat", padx=15, pady=8,
                 command=self._log_temizle).pack(side="right")

        # Durum
        self.durum = tk.Label(ana, text="⏸ Bekleniyor",
                             font=("Segoe UI", 10), bg=BG, fg=MUTED)
        self.durum.pack(anchor="w")

        # Log
        tk.Label(ana, text="📋 İşlem Geçmişi", font=("Segoe UI", 10, "bold"),
                bg=BG, fg=TEXT).pack(anchor="w", pady=(10, 3))
        self.log = scrolledtext.ScrolledText(ana, height=12, font=("Consolas", 9),
                                            bg=SURFACE, fg=TEXT, insertbackground=TEXT,
                                            relief="flat", bd=5, state="disabled")
        self.log.pack(fill="both", expand=True, pady=(0, 15))

        # Renk etiketleri
        self.log.tag_config("basari", foreground=GREEN)
        self.log.tag_config("hata", foreground=RED)
        self.log.tag_config("bilgi", foreground="#60a5fa")
        self.log.tag_config("uyari", foreground="#fbbf24")

        if not PDF_DESTEKLI:
            self._log("⚠ PDF desteği için: pip install PyPDF2", "uyari")
        self._log("Uygulama hazır. Klasör seçip izlemeyi başlatın.", "bilgi")

    def _klasor_sec(self):
        klasor = filedialog.askdirectory(title="İzlenecek klasörü seç")
        if klasor:
            self.klasor_yolu.set(klasor)
            self._log(f"📁 Klasör seçildi: {klasor}", "bilgi")

    def _izleme_toggle(self):
        if not self.izleniyor:
            self._izlemeyi_baslat()
        else:
            self._izlemeyi_durdur()

    def _izlemeyi_baslat(self):
        klasor = self.klasor_yolu.get()
        api_key = self.api_key.get()

        if not klasor:
            messagebox.showwarning("Uyarı", "Lütfen bir klasör seçin!")
            return
        if not api_key:
            messagebox.showwarning("Uyarı", "Lütfen API anahtarını girin!")
            return

        handler = DosyaHandler(self)
        self.observer = Observer()
        self.observer.schedule(handler, klasor, recursive=False)
        self.observer.start()
        self.izleniyor = True

        self.btn_izle.config(text="⏹ İzlemeyi Durdur", bg="#f87171")
        self.durum.config(text=f"🟢 İzleniyor: {klasor}", fg="#4ade80")
        self._log(f"🟢 İzleme başladı: {klasor}", "basari")
        self._log("📄 .docx veya .pdf dosyası atın, otomatik işlenecek!", "bilgi")

    def _izlemeyi_durdur(self):
        if self.observer:
            self.observer.stop()
            self.observer.join()
            self.observer = None
        self.izleniyor = False
        self.btn_izle.config(text="▶ İzlemeyi Başlat", bg="#4ade80")
        self.durum.config(text="⏸ Bekleniyor", fg="#6b6b80")
        self._log("⏹ İzleme durduruldu.", "uyari")

    def _manuel_isle(self):
        dosyalar = filedialog.askopenfilenames(
            title="Dosya seç",
            filetypes=[("Word/PDF", "*.docx *.pdf"), ("Word", "*.docx"), ("PDF", "*.pdf")]
        )
        for dosya in dosyalar:
            self.dosya_isle(dosya)

    def dosya_isle(self, dosya_yolu: str):
        def isle():
            api_key = self.api_key.get()
            istek = self.istek.get() or "tüm veriler"

            if not api_key:
                self._log("❌ API anahtarı eksik!", "hata")
                return

            yol = Path(dosya_yolu)
            self._log(f"⏳ İşleniyor: {yol.name}", "bilgi")

            try:
                metin = dosya_oku(dosya_yolu)
                if not metin:
                    self._log(f"❌ Metin okunamadı: {yol.name}", "hata")
                    return

                self._log(f"   📖 {len(metin)} karakter okundu", "bilgi")
                veri = groq_ile_cikart(metin, istek, api_key)

                zaman = datetime.now().strftime("%H%M%S")
                cikti = yol.parent / f"{yol.stem}_cikti_{zaman}.xlsx"
                satirlar, sutunlar = excel_olustur(veri, str(cikti))

                if satirlar == 0:
                    self._log(f"⚠ Veri bulunamadı: {yol.name}", "uyari")
                else:
                    self._log(f"✅ {yol.name} → {cikti.name}", "basari")
                    self._log(f"   📊 {satirlar} satır, {sutunlar} sütun", "basari")

            except Exception as e:
                self._log(f"❌ Hata: {e}", "hata")

        threading.Thread(target=isle, daemon=True).start()

    def _log(self, mesaj: str, tur: str = ""):
        def yaz():
            self.log.config(state="normal")
            zaman = datetime.now().strftime("%H:%M:%S")
            self.log.insert("end", f"[{zaman}] {mesaj}\n", tur)
            self.log.see("end")
            self.log.config(state="disabled")
        self.root.after(0, yaz)

    def _log_temizle(self):
        self.log.config(state="normal")
        self.log.delete("1.0", "end")
        self.log.config(state="disabled")

    def kapat(self):
        if self.observer:
            self.observer.stop()
            self.observer.join()
        self.root.destroy()


# ─── Başlat ──────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    root = tk.Tk()
    app = Uygulama(root)
    root.protocol("WM_DELETE_WINDOW", app.kapat)
    root.mainloop()
