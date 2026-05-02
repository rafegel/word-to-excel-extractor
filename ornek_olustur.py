"""Test amaçlı örnek Word dosyası oluşturur."""
from docx import Document

doc = Document()
doc.add_heading("Şirket Çalışan Raporu - 2024", 0)
doc.add_paragraph("Bu rapor şirketimizin aktif çalışanlarına ait bilgileri içermektedir.")

doc.add_heading("Yazılım Departmanı", 1)
tablo = doc.add_table(rows=1, cols=4)
tablo.style = "Table Grid"
baslik = tablo.rows[0].cells
baslik[0].text = "Ad Soyad"
baslik[1].text = "Pozisyon"
baslik[2].text = "Maaş (TL)"
baslik[3].text = "E-posta"

veriler = [
    ("Ahmet Yılmaz", "Kıdemli Yazılım Geliştirici", "45.000", "ahmet@sirket.com"),
    ("Zeynep Kaya", "Backend Developer", "38.000", "zeynep@sirket.com"),
    ("Mert Demir", "Frontend Developer", "35.000", "mert@sirket.com"),
    ("Selin Arslan", "DevOps Mühendisi", "42.000", "selin@sirket.com"),
]
for veri in veriler:
    satir = tablo.add_row().cells
    for i, deger in enumerate(veri):
        satir[i].text = deger

doc.add_heading("Pazarlama Departmanı", 1)
tablo2 = doc.add_table(rows=1, cols=4)
tablo2.style = "Table Grid"
baslik2 = tablo2.rows[0].cells
baslik2[0].text = "Ad Soyad"
baslik2[1].text = "Pozisyon"
baslik2[2].text = "Maaş (TL)"
baslik2[3].text = "E-posta"

veriler2 = [
    ("Fatma Çelik", "Pazarlama Müdürü", "55.000", "fatma@sirket.com"),
    ("Can Öztürk", "Sosyal Medya Uzmanı", "28.000", "can@sirket.com"),
    ("Ayşe Şahin", "İçerik Stratejisti", "32.000", "ayse@sirket.com"),
]
for veri in veriler2:
    satir = tablo2.add_row().cells
    for i, deger in enumerate(veri):
        satir[i].text = deger

doc.add_heading("Genel Notlar", 1)
doc.add_paragraph("Maaşlar brüt olarak belirtilmiştir. Performans değerlendirmesi yılda iki kez yapılmaktadır.")
doc.add_paragraph("İletişim: ik@sirket.com")

doc.save("ornek_calisanlar.docx")
print("✅ ornek_calisanlar.docx oluşturuldu.")
